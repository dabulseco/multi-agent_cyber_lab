"""Helpers that rebuild the lab manuals in the exact visual language of the v1.0
documents: Arial throughout, 1F4E79/2E75B6 heading blues, 2E75B6 table headers with
white text, BDD7EE zebra rows, FFF2CC callouts, DEEAF1 terminal blocks.

Each document is regenerated from its own v1.0 file used as a template, so
styles.xml, numbering.xml (numId 2 = bullet, 3 = decimal), the running header and
the section/page setup all carry over untouched.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
INK = "1A1A1A"
DEEP_BLUE = "1F4E79"
BLUE = "2E75B6"
ZEBRA = "BDD7EE"
WHITE = "FFFFFF"
TERMINAL_FILL = "DEEAF1"
CALLOUT_FILL = "FFF2CC"
TABLE_W = 9360
CELL_BORDER = "CCCCCC"


def _sz(el, size):
    for tag in ("w:sz", "w:szCs"):
        e = OxmlElement(tag)
        e.set(qn("w:val"), str(size))
        el.append(e)


def _rpr(run, *, size=22, color=INK, bold=False, italic=False):
    rpr = run._r.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:cs", "w:eastAsia", "w:hAnsi"):
        fonts.set(qn(a), FONT)
    rpr.append(fonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))
    if italic:
        rpr.append(OxmlElement("w:i"))
        rpr.append(OxmlElement("w:iCs"))
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    _sz(rpr, size)


def _spacing(par, before, after):
    ppr = par._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    ppr.append(sp)


def _center(par):
    ppr = par._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)


def _pstyle(par, style_id):
    """Set pStyle by styleId. The template defines duplicate style names, which
    breaks python-docx's name-based `add_paragraph(style=...)` lookup."""
    ppr = par._p.get_or_add_pPr()
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    ppr.append(ps)


def _numbered(par, num_id):
    ppr = par._p.get_or_add_pPr()
    npr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    npr.append(ilvl)
    npr.append(nid)
    ppr.append(npr)


def wipe_body(doc):
    """Drop every paragraph and table, keeping the trailing sectPr (page setup,
    header reference) so the regenerated document keeps its layout."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


# ---------------------------------------------------------------- block writers

def title_block(doc, title, subtitle, tagline, product, version):
    p = doc.add_paragraph()
    _spacing(p, 0, 240)
    _center(p)
    _rpr(p.add_run(title), size=56, color=DEEP_BLUE, bold=True)

    p = doc.add_paragraph()
    _spacing(p, 0, 160)
    _center(p)
    _rpr(p.add_run(subtitle), size=32, color=BLUE)

    p = doc.add_paragraph()
    _spacing(p, 160, 80)
    _center(p)
    _rpr(p.add_run(tagline), size=22, italic=True)

    for text in (product, version):
        p = doc.add_paragraph()
        _spacing(p, 0, 80)
        _center(p)
        _rpr(p.add_run(text), size=22)


def h1(doc, text):
    p = doc.add_paragraph()
    _pstyle(p, "Heading1")
    _spacing(p, 320, 160)
    _rpr(p.add_run(text), size=32, color=DEEP_BLUE, bold=True)


def h2(doc, text):
    p = doc.add_paragraph()
    _pstyle(p, "Heading2")
    _spacing(p, 240, 120)
    _rpr(p.add_run(text), size=26, color=BLUE, bold=True)


def h3(doc, text):
    p = doc.add_paragraph()
    _pstyle(p, "Heading3")
    _spacing(p, 200, 80)
    _rpr(p.add_run(text), size=24, color=INK, bold=True)


def body(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 60, 60)
    _rpr(p.add_run(text), size=22)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        _pstyle(p, "ListParagraph")
        _numbered(p, 2)
        _spacing(p, 40, 40)
        _rpr(p.add_run(item), size=22)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph()
        _pstyle(p, "ListParagraph")
        _numbered(p, 3)
        _spacing(p, 40, 40)
        _rpr(p.add_run(item), size=22)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcpr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        b = OxmlElement(side)
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), CELL_BORDER)
        b.set(qn("w:sz"), "1")
        borders.append(b)
    tcpr.append(borders)
    mar = OxmlElement("w:tcMar")
    for side, w in (("w:top", 120), ("w:left", 160), ("w:bottom", 120), ("w:right", 160)):
        m = OxmlElement(side)
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), str(w))
        mar.append(m)
    tcpr.append(mar)


def _cell_text(cell, text, *, size=22, color=INK, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    _spacing(p, 40, 40)
    lines = text.split("\n")
    _rpr(p.add_run(lines[0]), size=size, color=color, bold=bold)
    for extra in lines[1:]:
        q = cell.add_paragraph()
        _spacing(q, 40, 40)
        _rpr(q.add_run(extra), size=size, color=color, bold=bold)


def _table_shell(doc, cols):
    t = doc.add_table(rows=0, cols=cols)
    tblpr = t._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:type"), "dxa")
    w.set(qn("w:w"), str(TABLE_W))
    tblpr.append(w)
    borders = OxmlElement("w:tblBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        b = OxmlElement(side)
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), "auto")
        b.set(qn("w:sz"), "4")
        borders.append(b)
    tblpr.append(borders)
    return t


def table(doc, header, rows):
    """Header row in solid blue with white bold text, body rows zebra-striped."""
    t = _table_shell(doc, len(header))
    hr = t.add_row()
    for cell, text in zip(hr.cells, header):
        _shade(cell, BLUE)
        _cell_text(cell, text, color=WHITE, bold=True)
    for i, row in enumerate(rows):
        r = t.add_row()
        fill = WHITE if i % 2 == 0 else ZEBRA
        for cell, text in zip(r.cells, row):
            _shade(cell, fill)
            _cell_text(cell, text)
    spacer(doc)
    return t


def callout(doc, label, text):
    t = _table_shell(doc, 1)
    cell = t.add_row().cells[0]
    _shade(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    _spacing(p, 40, 40)
    _rpr(p.add_run(f"{label}  "), size=22, color=DEEP_BLUE, bold=True)
    _rpr(p.add_run(text), size=22)
    spacer(doc)


def terminal(doc, command):
    t = _table_shell(doc, 1)
    cell = t.add_row().cells[0]
    _shade(cell, TERMINAL_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    _spacing(p, 40, 40)
    _rpr(p.add_run("Terminal:  "), size=22, color=DEEP_BLUE, bold=True)
    _rpr(p.add_run(command), size=22)
    spacer(doc)


def spacer(doc):
    p = doc.add_paragraph()
    _spacing(p, 0, 60)


def open_template(path):
    doc = Document(path)
    wipe_body(doc)
    return doc
